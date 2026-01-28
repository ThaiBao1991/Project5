import certifi
import requests
# ca_bundle_path = certifi.where()  # Get the path to the certifi CA bundle
# requests.get('https://cdn.extendoffice.com', verify=ca_bundle_path)
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches
from urllib.parse import urljoin
import urllib3
from urllib3.exceptions import InsecureRequestWarning

def download_and_add_image(doc, img_url, save_path):
    """
    Tải xuống hình ảnh và thêm vào tài liệu Word.

    Args:
        doc: Đối tượng Document của Word.
        img_url: URL của hình ảnh.
        save_path: Đường dẫn lưu trữ hình ảnh tạm thời.
    """
    try:
        # response = requests.get(img_url, stream=True)
        # response = requests.get(img_url, verify=False)
        print(img_url)
        ca_bundle_path = certifi.where()
        response = requests.get(img_url, verify=ca_bundle_path)
        response.raise_for_status()

        with open(save_path, 'wb') as out_file:
            for chunk in response.iter_content(1024):
                out_file.write(chunk)
        doc.add_picture(save_path)
    except requests.exceptions.RequestException as e:
        print(f"Error downloading image from {img_url}: {e}")

def get_absolute_url(base_url, relative_url):
    """
    Hàm tạo URL tuyệt đối từ URL gốc và URL tương đối.

    Args:
        base_url (str): URL gốc của trang web.
        relative_url (str): URL tương đối.

    Returns:
        str: URL tuyệt đối.
    """
    return urljoin(base_url, relative_url)
base_url = 'https://vi.extendoffice.com'

# Ham lay du lieu ve file doc
def recursive_parse(document, ele, parent_paragraph=None):
    document.styles['Normal'].font.name = 'Times New Roman'
    if ele.name is not None:
        for child in ele.children:
            if child.name =="img":
                print (child)
            if child.name == 'p':
                if parent_paragraph is None:
                    parent_paragraph = document.add_paragraph(child.text)
                else:
                    parent_paragraph.add_run(child.text)
                for sub_child in child.children:
                    recursive_parse(document, sub_child, parent_paragraph)
            elif child.name in ['h1', 'h2']:
                level = int(child.name[-1]) - 1
                parent_paragraph = document.add_paragraph(child.text, style=f'Heading {level}')
                for sub_child in child.children:
                    recursive_parse(document, sub_child, parent_paragraph)
            elif child.name == 'a':
                if parent_paragraph is None:
                    parent_paragraph = document.add_paragraph()
                run = parent_paragraph.add_run(child.text)
                # Áp dụng định dạng cho liên kết nếu cần
            elif child.name == 'img':
                img_url = child.get('src')
                if img_url:
                    save_path = 'temp_image.jpg'
                    download_and_add_image(document, img_url, save_path)
            elif child.name == 'div':
                recursive_parse(document, child, parent_paragraph)
            else:
                recursive_parse(document, child, parent_paragraph)

def download_and_add_image(document, img_url, save_path):
    import requests
    from PIL import Image
    from io import BytesIO

    response = requests.get(img_url)
    if response.status_code == 200:
        image = Image.open(BytesIO(response.content))
        image.save(save_path)
        document.add_picture(save_path)



def create_word_document(url):
    # Gửi yêu cầu HTTP và phân tích HTML
    response = requests.get(url)
    soup = BeautifulSoup(response.content, 'html.parser')

    # Tìm tất cả các thẻ <a> bên trong ul có id="ul-search"
    links = soup.find('ul', id='ul-search').find_all('a')

    # Tạo file Word mới
    doc = Document()

    # Tạo mục lục
    doc.add_heading('Mục lục', level=0)

    # Xử lý lỗi lấy nội dung
    for link in links:
        text = link.text.strip()
        href = link['href']

        # Thêm vào mục lục
        doc.add_paragraph(text, style='Heading 2')

        # In ra màn hình để kiểm tra (tùy chọn)
        
        # print(f"Link: {href}")
        # print(f"Text: {text}")
        absolute_url = get_absolute_url(base_url, href)
        try:
            response = requests.get(absolute_url)
            response.raise_for_status()
            soup = BeautifulSoup(response.content, 'html.parser')

            # Tìm thẻ div theo class
            target_div = soup.find('div', class_='uk-margin-small-top')
            
            recursive_parse(doc,target_div)

            # Lưu file Word
            doc.save('output.docx')
        except requests.exceptions.RequestException as e:
            print(f"Error fetching content from {href}: {e}")
        doc.save('output.docx') 
    # Lưu file Word
    doc.save('output.docx') 

# URL của trang web
url = "https://vi.extendoffice.com/documents/excel"  # Thay thế bằng URL thực tế

create_word_document(url)
