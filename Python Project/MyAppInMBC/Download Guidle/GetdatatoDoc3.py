import os
import time
import pyautogui
import pywinauto
import requests
import sys
from docx import Document
from urllib.parse import urljoin
from bs4 import BeautifulSoup
from tkinter import Tk, filedialog, messagebox
from pathlib import Path
from pywinauto.findwindows import ElementNotFoundError
from PIL import ImageGrab, ImageChops, Image
from io import BytesIO
import win32com.client
from pywinauto.application import Application
import pyperclip
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import TimeoutException
import base64
import urllib3

urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

user_dir = Path("C:/Users/12953 bao/Desktop/desktop/work/Project/Python/BasicLearnPython/W3schools")
output_path = user_dir / "output.docx"
base_url = 'https://vi.extendoffice.com'

def convert_webp_to_png(webp_path, png_path):
    try:
        img = Image.open(webp_path)
        img.save(png_path, "PNG")
        print(f"Đã chuyển đổi ảnh: {png_path}")
    except Exception as e:
        print(f"Lỗi khi chuyển đổi ảnh: {e}")

def image_to_base64(image_path):
    with open(image_path, "rb") as image_file:
        return base64.b64encode(image_file.read()).decode("utf-8")

def select_image_file(default_path):
    root = Tk()
    root.withdraw()
    root.attributes('-topmost', True)
    file_path = filedialog.askopenfilename(initialdir=os.path.dirname(default_path), title="Select Image File",
                                           filetypes=(("PNG files", "*.png"), ("All files", "*.*")))
    root.destroy()
    return file_path

def select_word_file():
    root = Tk()
    root.withdraw()
    root.attributes('-topmost', True)
    file_path = filedialog.askopenfilename(title="Select Word File",
                                           filetypes=(("Word files", "*.docx"), ("All files", "*.*")))
    root.destroy()
    return file_path

def has_window_changed(dlg, initial_screenshot):
    current_screenshot = capture_screenshot(dlg)
    return current_screenshot != initial_screenshot

def open_file_dialog(dlg):
    initial_screenshot = capture_screenshot(dlg)
    max_attempts = 5
    attempt = 0

    while attempt < max_attempts:
        dlg.set_focus()
        dlg.type_keys('^o')
        time.sleep(5)

        if has_window_changed(dlg, initial_screenshot):
            print("Cửa sổ đã thay đổi. Tiếp tục thực thi...")
            return True

        attempt += 1
        print(f"Thử lần {attempt} không thành công, thử lại...")

    print("Mở file chờ 15s chưa có tín hiệu.")
    return False

def capture_screenshot(dlg):
    rect = dlg.rectangle()
    screenshot = ImageGrab.grab(bbox=(rect.left, rect.top, rect.right, rect.bottom))
    return screenshot

def has_window_changed(dlg, initial_screenshot):
    current_screenshot = capture_screenshot(dlg)
    diff = ImageChops.difference(initial_screenshot, current_screenshot)
    return diff.getbbox() is not None

def wait_for_window(app, title_re, timeout=10):
    start_time = time.time()
    while time.time() - start_time < timeout:
        try:
            dlg = app.window(title_re=title_re)
            if dlg.exists(timeout=1):
                return dlg
        except ElementNotFoundError:
            pass
        time.sleep(0.5)
    return None

def wait_for_save_completion(file_path, timeout=30):
    start_time = time.time()
    file_size = -1

    while time.time() - start_time < timeout:
        try:
            if os.path.exists(file_path):
                current_size = os.path.getsize(file_path)
                if current_size == file_size:
                    return True
                else:
                    file_size = current_size
                    time.sleep(1)
            else:
                print("File không tồn tại.")
                return False
        except Exception as e:
            print(f"Lỗi trong quá trình kiểm tra: {e}")
            return False
    return False

def close_word_document(dlg):
    try:
        dlg.close()
        print("Đã đóng cửa sổ Word.")
    except pywinauto.application.findwindows.ElementNotFoundError:
        print("Cửa sổ Word đã được đóng trước đó hoặc không tìm thấy.")
    except Exception as e:
        print(f"Lỗi khi đóng cửa sổ Word: {e}")

def check_doc_is_open(document_name):
    try:
        word_app = win32com.client.GetObject(Class="Word.Application")
    except win32com.client.com_error:
        return None
    for doc in word_app.Documents:
        if doc.Name == document_name:
            print("doc is : ", doc)
            return doc
        else:
            return word_app.ActiveDocument

def get_word_count(document_name=None):
    try:
        word_app = win32com.client.GetObject(Class="Word.Application")
        doc = check_doc_is_open(document_name)
        word_count = doc.Words.Count
        return word_count
    except AttributeError as e:
        print(f"Lỗi thuộc tính: {e}")
        return None
    except Exception as e:
        print(f"Lỗi không xác định: {e}")
        return None

def wait_for_paste_completion(initial_word_count, document_name, timeout=30):
    start_time = time.time()
    while time.time() - start_time < timeout:
        current_word_count = get_word_count(document_name)
        if current_word_count > initial_word_count:
            return True
        time.sleep(1)
    return False

def find_element_with_timeout(driver, locator, timeout=10):
    try:
        element = WebDriverWait(driver, timeout).until(EC.presence_of_element_located(locator))
        return element
    except TimeoutException:
        return None
    except Exception as e:
        print(f"Lỗi không xác định: {e}")
        return None

def process_element(driver, element):
    image_dir = os.path.join("project", "imagetmp")
    os.makedirs(image_dir, exist_ok=True)

    js_get_webp_images = """
    var images = document.querySelectorAll('img');
    var webpUrls = [];
    images.forEach(img => {
        if (img.src.endsWith('.webp')) {
            webpUrls.push(img.src);
        }
    });
    return webpUrls;
    """
    webp_urls = driver.execute_script(js_get_webp_images)

    for i, url in enumerate(webp_urls):
        try:
            absolute_url = get_absolute_url(base_url, url)
            print("Đang tải ảnh từ URL:", absolute_url)

            response = requests.get(absolute_url, verify=False)
            if response.status_code == 200:
                webp_path = os.path.join(image_dir, f"image_{i}.webp")
                with open(webp_path, "wb") as f:
                    f.write(response.content)
                print(f"Đã tải ảnh: {webp_path}")

                png_path = os.path.join(image_dir, f"image_{i}.png")
                convert_webp_to_png(webp_path, png_path)

                os.remove(webp_path)
                print(f"Đã xóa file .webp: {webp_path}")

                png_base64 = image_to_base64(png_path)
                js_replace_src = f"""
                var images = document.querySelectorAll('img');
                images.forEach(img => {{
                    if (img.src === "{absolute_url}") {{
                        img.src = "data:image/png;base64,{png_base64}";
                    }}
                }});
                """
                driver.execute_script(js_replace_src)
                print(f"Đã thay thế src của ảnh: {absolute_url}")
            else:
                print(f"Không thể tải ảnh từ URL: {absolute_url}")
        except Exception as e:
            print(f"Lỗi khi tải ảnh: {e}")

    js_remove_elements = """
    var elements = document.querySelectorAll('.uk-margin-remove-last-child.custom');
    elements.forEach(element => {
        var firstChild = element.firstElementChild;
        if (firstChild && firstChild.tagName === 'H3' && firstChild.innerText.trim() === 'Công cụ năng suất văn phòng tốt nhất') {
            element.parentNode.removeChild(element);
        }
    });
    """
    driver.execute_script(js_remove_elements)
    print("Đã xóa các đối tượng thỏa mãn điều kiện.")

    for filename in os.listdir(image_dir):
        if filename.endswith(".png"):
            file_path = os.path.join(image_dir, filename)
            os.remove(file_path)
            print(f"Đã xóa file .png: {file_path}")

    if not element:
        print("Không có element nào để xử lý.")
        return

    driver.execute_script("arguments[0].scrollIntoView();", element)
    driver.execute_script("arguments[0].focus();", element)
    time.sleep(1)

    js_code = """
    var element = arguments[0];
    var stopElement = document.querySelector('.uk-margin-remove-last-child.custom h3[style="margin-top: ' + arguments[1] + 'px;"]');
    var range = document.createRange();

    if (stopElement) {
        range.setStartBefore(element);
        range.setEndBefore(stopElement);
    } else {
        range.setStartBefore(element);
        range.setEndAfter(document.body.lastChild);
    }

    var divs = document.querySelectorAll('.uk-margin-remove-last-child.custom');

    divs.forEach(function(div) {
        if (div.querySelector('style')) {
            div.parentNode.removeChild(div);
        }
    });

    var sel = window.getSelection();
    sel.removeAllRanges();
    sel.addRange(range);

    var messageBox = document.createElement('div');
    messageBox.style.position = 'fixed';
    messageBox.style.top = '10px';
    messageBox.style.left = '50%';
    messageBox.style.transform = 'translateX(-50%)';
    messageBox.style.padding = '10px';
    messageBox.style.backgroundColor = 'lightgreen';
    messageBox.style.border = '1px solid green';
    messageBox.style.zIndex = '10000';
    messageBox.innerText = 'Đã chọn đối tượng!';
    document.body.appendChild(messageBox);

    setTimeout(function() {
        document.body.removeChild(messageBox);
    }, 2000);
    """
    driver.execute_script(js_code, element)

    pyautogui.hotkey('ctrl', 'c')
    time.sleep(1)

def find_element_by_multiple_locators(driver, locators):
    for by, value in locators:
        try:
            element = driver.find_element(by, value)
            if element:
                return element
        except:
            continue
    return None

def copy_element_content(driver, element):
    if not element:
        print("Không có element nào để xử lý.")
        return

    driver.execute_script("""
        const element = arguments[0];
        const range = document.createRange();
        range.selectNodeContents(element);
        const selection = window.getSelection();
        selection.removeAllRanges();
        selection.addRange(range);
        document.execCommand('copy');
        selection.removeAllRanges();
    """, element)

def copy_and_paste_content(driver, document, output_path):
    locators = [
        (By.CLASS_NAME, "uk-margin-small-top"),
        (By.CSS_SELECTOR, ".uk-width-expand\\@m.uk-first-column"),
        (By.CSS_SELECTOR, ".some-other-class"),
        (By.ID, "some-id")
    ]

    try:
        element = find_element_by_multiple_locators(driver, locators)
        print(element)
        if element:
            process_element(driver, element)
            time.sleep(1)
            pyautogui.hotkey('ctrl', 'c')
        else:
            print("Không tìm thấy element nào với các locator đã cho.")
            driver.execute_script("""
            var messageBox = document.createElement('div');
            messageBox.style.position = 'fixed';
            messageBox.style.top = '10px';
            messageBox.style.left = '50%';
            messageBox.style.transform = 'translateX(-50%)';
            messageBox.style.padding = '10px';
            messageBox.style.backgroundColor = 'lightcoral';
            messageBox.style.border = '1px solid red';
            messageBox.style.zIndex = '10000';
            messageBox.innerText = 'Không tìm thấy đối tượng!';
            document.body.appendChild(messageBox);

            setTimeout(function() {
                document.body.removeChild(messageBox);
            }, 2000);
            """)
    except Exception as e:
        print(f"Lỗi tổng quát: {e}")

    path1 = "C:\\Program Files (x86)\\Microsoft Office\\root\\Office16\\winword.exe"
    path2 = "C:\\Program Files\\Microsoft Office\\root\\Office16\\winword.exe"
    if os.path.exists(path1):
        app_path = path1
    elif os.path.exists(path2):
        app_path = path2
    else:
        raise FileNotFoundError("Không tìm thấy Microsoft Word ở bất kỳ đường dẫn nào.")
    app = pywinauto.Application().start(app_path)

    if not app.windows():
        app.start(app_path)

    try:
        dlg = app.window(title_re=".*Word.*")
        dlg.wait('visible', timeout=40)

        if open_file_dialog(dlg):
            pass
        else:
            print("Không thể mở cửa sổ Open.")

        image_path = 'Python Tutorial\\browse_button_image.png'

        if not os.path.exists(image_path):
            print(f"Không tìm thấy file ảnh tại {image_path}. Vui lòng chọn file ảnh mới.")
            image_path = select_image_file(image_path)
            if not image_path:
                print("Không có file ảnh nào được chọn.")
                return

        browse_button_location = None
        for i in range(10):
            browse_button_location = pyautogui.locateCenterOnScreen(image_path, confidence=0.8)
            if browse_button_location:
                print(f"Đã tìm thấy file ảnh tại {image_path}.")
                break
            time.sleep(1)

        if not browse_button_location:
            print(f"Không tìm thấy file ảnh tại {image_path}. Vui lòng chọn file ảnh mới.")
            new_image_path = select_image_file(image_path)
            if new_image_path:
                image_path = new_image_path
                browse_button_location = pyautogui.locateCenterOnScreen(image_path, confidence=0.8)
                if browse_button_location:
                    print(f"Đã tìm thấy file ảnh tại {image_path}.")
                else:
                    print("Không tìm thấy nút 'Browse' trong file ảnh mới.")
                    return
            else:
                print("Không có file ảnh nào được chọn.")
                return

        print(f"Moving to: {browse_button_location}")
        pyautogui.moveTo(browse_button_location)
        time.sleep(0.5)
        pyautogui.click(browse_button_location)

        print(output_path)

        directory_path = os.path.dirname(output_path)
        file_name = os.path.basename(output_path)
        new_patch = Path(directory_path) / file_name

        dlg_open = app.window(title_re=".*Open.*")

        start_time = time.time()
        dlg_open.wait('ready', timeout=1)
        end_time = time.time()
        print(f"Thời gian chờ thực tế bảng open ready: {end_time - start_time} giây")

        dlg_open.type_keys(str(new_patch), with_spaces=True, pause=0)
        time.sleep(0.1)
        dlg_open.type_keys('{ENTER}')
        time.sleep(0.3)

        find_doc_name = ".*" + file_name + ".*"
        print("find_doc_name là :", find_doc_name)

        dlg_word_open = None
        try:
            dlg_word_open = app.window(title_re=find_doc_name)

            max_attempts = 6
            attempt = 0
            while attempt < max_attempts:
                try:
                    start_time = time.time()
                    dlg_word_open.wait('ready', timeout=1.5)
                    end_time = time.time()
                    print(f"Thời gian chờ thực tế bảng dlg_word_open ready: {end_time - start_time} giây")
                    break
                except Exception as e:
                    attempt += 1
                    print(f"Thử lần {attempt} không thành công, thử lại...")

            print("File Word đã mở xong.")
            time.sleep(0.5)
            dlg_word_open.type_keys('^{END}')
            dlg_word_open.type_keys('{ENTER}')
        except pywinauto.findwindows.ElementNotFoundError:
            print("Không tìm thấy cửa sổ Word. Có thể có lỗi khi mở file.")
        except Exception as e:
            print(f"Lỗi: {e}")

        initial_word_count = get_word_count(file_name)

        print(f"so ky tu den duoc là initial_word_count := '{initial_word_count}'")
        print("so ky tu den duoc là initial_word_count := ", initial_word_count)

        doc_check = check_doc_is_open(output_path)
        words_before = 0
        words_after = 0
        if doc_check:
            try:
                words_before = doc_check.Words.Count
                print(f"Number of words before: {words_before}")
            except Exception as e:
                print(f"An error occurred while counting words: {e}")
            else:
                print(f"Document '{output_path}' not found open in Word.")

        pyautogui.hotkey('ctrl', 'v')
        start_time = time.time()
        timeout = 5 * 60

        while True:
            try:
                words_after = doc_check.Words.Count
                print(f"Number of words after: {words_after}")
            except Exception as e:
                print(f"An error occurred while counting words: {e}")

            if words_after > words_before:
                break
            else:
                print("Không có sự thay đổi khi dán")
            elapsed_time = time.time() - start_time
            if elapsed_time > timeout:
                print(f"Đã quá thời gian chờ ({timeout} giây). Dừng kiểm tra.")
                break

            time.sleep(1)

        pyautogui.hotkey('ctrl', 's')

        if wait_for_save_completion(output_path, timeout=30):
            print("Lưu file hoàn thành.")
        else:
            print("Lưu file không hoàn thành trong thời gian chờ.")

        dlg_word_open = app.window(title_re=find_doc_name)
        close_word_document(dlg_word_open)
    except pywinauto.findwindows.ElementNotFoundError:
        print("Không tìm thấy cửa sổ Word")

def get_absolute_url(base_url, relative_url):
    return urljoin(base_url, relative_url)

def get_existing_links(document):
    existing_links = []
    for paragraph in document.paragraphs:
        if paragraph.style.name == 'Heading 2':
            existing_links.append(paragraph.text.strip())
    return existing_links

def create_word_document(url, document, output_path):
    response = requests.get(url)
    soup = BeautifulSoup(response.content, 'html.parser')
    links = soup.find('ul', id='ul-search').find_all('a')

    for i, link in enumerate(links):
        if i < 20:
            print(f"Link {i}: {link.text.strip()} - {link['href']}")

    existing_links = get_existing_links(document)
    with open('links.txt', 'w', encoding='utf-8') as file:
        for i, link in enumerate(links):
            text = link.text.strip()
            href = link['href']
            file.write(f'Link {i}: {text} - {href}\n')
            print(f"Link la : '{text}'")

            if text in existing_links:
                print(f"Link '{text}' đã tồn tại trong tài liệu.")
                continue

            absolute_url = get_absolute_url(base_url, href)
            print(f"Đang truy cập link {i}: {absolute_url}")
            driver = webdriver.Chrome()
            driver.get(absolute_url)
            copy_and_paste_content(driver, document, output_path)
            driver.quit()
            print(f"Đã hoàn thành link {i}")

def main():
    root = Tk()
    root.withdraw()
    root.attributes('-topmost', True)
    response = messagebox.askyesno("Chọn tùy chọn", "Bạn có muốn tạo file mới không? (Chọn 'No' để cập nhật file cũ)")
    root.destroy()

    if response:
        document = Document()
        document.add_paragraph('')
        current_dir = os.getcwd()
        print("current dir là : ", current_dir)
        output_path = 'output.docx'
        output_path = os.path.join(current_dir, output_path)
        document.save(output_path)
    else:
        word_file_path = select_word_file()
        if word_file_path:
            document = Document(word_file_path)
            output_path = word_file_path
        else:
            print("Không có file nào được chọn.")
            return

    url = "https://vi.extendoffice.com/documents/excel"
    create_word_document(url, document, output_path)

if __name__ == "__main__":
    main()